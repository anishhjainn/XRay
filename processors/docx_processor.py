# processors/docx_processor.py
"""
DOCX technician (FileProcessor): builds a read-only FileArtifact for .docx files.

Existing metadata produced (consumed by current checks):
- comments_present: bool
- comments_count: int
- tracked_changes_count: int
- highlight_run_count: int          (w:highlight)
- shading_highlight_count: int      (w:shd with non-empty/non-auto fill)
- core_modified: ISO 8601 string (from core properties)
- read_error / read_error_detail on core parse failures

New in Phase 3 (for spelling now, grammar later):
- text_sample: normalized plain text sample of document body + tables (read-only)
- text_length: length of text_sample
- token_count: basic English word token count (lowercased)
- text_extraction_error / text_extraction_error_detail: flags if text extraction fails

Design (SOLID):
- SRP: Processor gathers metadata only; it does not make policy decisions (checks do).
- OCP: Adding new checks only reads metadata; orchestrator is unchanged.
- LSP: Still honors FileProcessor; orchestrator treats it interchangeably.
- ISP: Checks depend only on FileArtifact (no parsing libraries).
- DIP: Checks depend on abstract text (string in metadata), not concrete libraries like python-docx.

Implementation notes:
- Use python-docx for core props (simple & stable).
- Use zipfile + iterparse to stream-read XML parts (fast, low memory) for existing counts.
- Use utils.text_extract.extract_docx_text to build a single text sample (early-stopped at max_text_chars from config).
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Optional
from zipfile import ZipFile, BadZipFile
from xml.etree.ElementTree import iterparse

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from core.interfaces import FileProcessor
from core.models import FileArtifact
from core.registry import register_processor

# NEW: text extraction helpers and config
from utils.text_extract import extract_docx_text, tokenize_words
from infra.config_loader import load_config


DOC_PARTS_TO_SCAN = (
    "word/document.xml",
    # headers / footers may contain highlights or tracked changes
    # We'll scan any present files matching these prefixes.
)


class DocxProcessor(FileProcessor):
    def supports(self):
        return [".docx"]

    def build_artifact(self, path: Path) -> FileArtifact:
        size = path.stat().st_size
        metadata: Dict[str, Any] = {
            "kind": "docx",
            "paragraph_count": None,
            "table_count": None,
            "core_author": None,
            "core_created": None,
            "core_modified": None,
            "comments_present": False,
            "comments_count": 0,
            "tracked_changes_count": 0,
            "highlight_run_count": 0,
            "shading_highlight_count": 0,
            "read_error": False,
            # NEW: text extraction fields
            "text_sample": "",
            "text_length": 0,
            "token_count": 0,
            # We keep separate flags for text extraction failures so other checks still run
            "text_extraction_error": False,
            # optional detail string set only when an error occurs
            # "text_extraction_error_detail": "...",
        }

        # 1) Core props via python-docx (graceful if it fails)
        try:
            doc = Document(str(path))
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            core = doc.core_properties
            metadata["core_author"] = getattr(core, "author", None)
            metadata["core_created"] = _safe_iso(getattr(core, "created", None))
            metadata["core_modified"] = _safe_iso(getattr(core, "modified", None))
        except Exception:
            # We'll still try XML scanning even if python-docx stumbles.
            # Keep already-initialized None values.
            pass

        # 2) XML scan inside the .docx (ZIP)
        try:
            with ZipFile(str(path)) as zf:
                # Comments (legacy + modern)
                c_count = 0
                # legacy comments: /word/comments.xml with <w:comment>
                if "word/comments.xml" in zf.namelist():
                    c_count += _count_tags_in_zip(zf, "word/comments.xml", {"comment"})
                # modern/extended comments: /word/commentsExtended.xml with w15:commentEx
                if "word/commentsExtended.xml" in zf.namelist():
                    c_count += _count_tags_in_zip(zf, "word/commentsExtended.xml", {"commentEx"})
                metadata["comments_count"] = c_count
                metadata["comments_present"] = c_count > 0

                # Tracked changes across main doc + headers/footers
                tracked_tags = {
                    "ins",
                    "del",
                    "moveFrom",
                    "moveTo",
                    "tblPrChange",
                    "trPrChange",
                    "tcPrChange",
                    "pPrChange",
                    "rPrChange",
                }
                tracked_total = 0

                # Main document
                if "word/document.xml" in zf.namelist():
                    tracked_total += _count_tags_in_zip(zf, "word/document.xml", tracked_tags)

                # Headers/footers (names vary: header1.xml, footer2.xml, etc.)
                for name in zf.namelist():
                    if name.startswith("word/header") and name.endswith(".xml"):
                        tracked_total += _count_tags_in_zip(zf, name, tracked_tags)
                    if name.startswith("word/footer") and name.endswith(".xml"):
                        tracked_total += _count_tags_in_zip(zf, name, tracked_tags)

                metadata["tracked_changes_count"] = tracked_total

                # Highlights: explicit w:highlight, plus shading w:shd fill
                highlight_count = 0
                shading_count = 0

                scan_targets = ["word/document.xml"]
                scan_targets += [n for n in zf.namelist() if n.startswith("word/header") and n.endswith(".xml")]
                scan_targets += [n for n in zf.namelist() if n.startswith("word/footer") and n.endswith(".xml")]

                for part in scan_targets:
                    h, s = _count_highlights_in_part(zf, part)
                    highlight_count += h
                    shading_count += s

                metadata["highlight_run_count"] = highlight_count
                metadata["shading_highlight_count"] = shading_count

        except (PackageNotFoundError, KeyError, ValueError, OSError, Exception) as exc:
            metadata["read_error"] = True
            metadata["read_error_detail"] = str(exc) or exc.__class__.__name__

        # 3) NEW: Plain-text sample extraction (for spelling/grammar checks)
        # We compute this even if earlier steps had issues, so downstream checks that only need text can still run.
        try:
            cfg = load_config()
            # Only extract text when at least one of these features is enabled.
            if cfg.get("enable_spelling", True) or cfg.get("enable_grammar", False):
                max_chars = int(cfg.get("max_text_chars", 5_000_000))
                sample = extract_docx_text(path, max_chars)
                metadata["text_sample"] = sample
                metadata["text_length"] = len(sample)
                # Tokenization gives a rough word count for quick stats; checks will re-tokenize as needed.
                metadata["token_count"] = len(tokenize_words(sample))
        except Exception as exc:
            # Important: text extraction failure should not mark the entire file unreadable.
            metadata["text_sample"] = ""
            metadata["text_length"] = 0
            metadata["token_count"] = 0
            metadata["text_extraction_error"] = True
            metadata["text_extraction_error_detail"] = str(exc) or exc.__class__.__name__

        return FileArtifact(
            path=path,
            extension=".docx",
            size_bytes=size,
            metadata=metadata,
        )


# --- helpers ---

def _safe_iso(dt) -> Optional[str]:
    try:
        return dt.isoformat() if dt is not None else None
    except Exception:
        return None


def _count_tags_in_zip(zf: ZipFile, member: str, localname_set: set[str]) -> int:
    """
    Count elements whose tag's local-name is in localname_set.
    local-name = tag without namespace, e.g., '{ns}comment' -> 'comment'.
    """
    count = 0
    with zf.open(member) as fp:
        # 'events' = ("end",) is cheaper; clear elements as we go to save memory
        for _event, elem in iterparse(fp, events=("end",)):
            if _local_name(elem.tag) in localname_set:
                count += 1
            elem.clear()
    return count


def _count_highlights_in_part(zf: ZipFile, member: str) -> tuple[int, int]:
    """
    Return (highlight_count, shading_count) for a single XML part.
    - highlight_count: number of <w:highlight> occurrences
    - shading_count: number of <w:shd> with a non-empty/non-'auto' fill
    """
    h_count = 0
    s_count = 0
    with zf.open(member) as fp:
        for _event, elem in iterparse(fp, events=("end",)):
            lname = _local_name(elem.tag)
            if lname == "highlight":
                h_count += 1
            elif lname == "shd":
                # Attributes may be namespaced; accept any attr ending with 'fill'
                has_fill = False
                for k, v in elem.attrib.items():
                    if k.endswith("fill"):
                        val = (v or "").strip().lower()
                        if val and val not in {"auto", "none"}:
                            has_fill = True
                            break
                if has_fill:
                    s_count += 1
            elem.clear()
    return h_count, s_count


def _local_name(tag: str) -> str:
    """
    Strip the XML namespace from a tag: '{namespace}name' -> 'name'.
    """
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


# Register on import
register_processor(DocxProcessor())